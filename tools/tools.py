from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD

from datetime import datetime
#import re


meses = (
    'janeiro',
    'fevereiro',
    'março',
    'abril',
    'maio',
    'junho',
    'julho',
    'agosto',
    'setembro',
    'outubro',
    'novembro',
    'dezembro'
)

def validar_cpf(cpf):
    """ Efetua a validação do CPF, tanto formatação quando dígito verificadores.

    Parâmetros:
        cpf (str): CPF a ser validado

    Retorno:
        bool:
            - Falso, quando o CPF não possuir o formato 999.999.999-99;
            - Falso, quando o CPF não possuir 11 caracteres numéricos;
            - Falso, quando os dígitos verificadores forem inválidos;
            - Verdadeiro, caso contrário.

    Exemplos:

    >>> validate('529.982.247-25')
    True
    >>> validate('52998224725')
    False
    >>> validate('111.111.111-11')
    False
    """

    # Verifica a formatação do CPF
#    if not re.match(r'\d{3}\.\d{3}\.\d{3}-\d{2}', cpf):
#        return False

    # Obtém apenas os números do CPF, ignorando pontuações
    numbers = [int(digit) for digit in cpf if digit.isdigit()]

    # Verifica se o CPF possui 11 números ou se todos são iguais:
    if len(numbers) != 11:
        return False

    # Validação do primeiro dígito verificador:
    sum_of_products = sum(a*b for a, b in zip(numbers[0:9], range(10, 1, -1)))
    expected_digit = (sum_of_products * 10 % 11) % 10
    if numbers[9] != expected_digit:
        return False

    # Validação do segundo dígito verificador:
    sum_of_products = sum(a*b for a, b in zip(numbers[0:10], range(11, 1, -1)))
    expected_digit = (sum_of_products * 10 % 11) % 10
    if numbers[10] != expected_digit:
        return False

    return True

def gerar_carta(cpf, nome, bolsa):
    doc = Document('Modelo_Proac.docx')

    data = str(datetime.today())
    data, _ = data.split(' ')
    ano, mes, dia = data.split('-')
    mes = meses[int(mes)-1]
    data = dia + ' de ' + mes + ' de ' + ano

    cpf_formatado = cpf[0:3] + '.' + cpf[3:6] + '.' + cpf[6:9] + '-' + cpf[-2:]
    nome = nome.upper()
    bolsa = bolsa.upper()

    doc.add_paragraph()
    doc.add_paragraph()
    para = doc.add_paragraph(f'Campos dos Goytacazes, {data}.')
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()
    para = doc.add_paragraph()
    bold_para = para.add_run('Ao Banco Bradesco S/A')
    bold_para.bold = True

    doc.add_paragraph()
    para = doc.add_paragraph()
    bold_para = para.add_run('Assunto: Abertura de Conta Corrente Pessoa Física')
    bold_para.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Senhor (a) Gerente,')

    doc.add_paragraph()
    para = doc.add_paragraph('A Universidade Estadual do Norte Fluminense Darcy Ribeiro - UENF, inscrita no CNPJ: 04.809.688/0001-06, encaminha por meio do presente documento, o(a) solicitante ')
    bold_para = para.add_run(f'{nome}')
    bold_para.bold = True
    bold_para = para.add_run(', portador(a) do CPF ')
    bold_para = para.add_run(f'{cpf_formatado}')
    bold_para.bold = True
    bold_para = para.add_run(f', para abertura da conta corrente TIPO PESSOA FÍSICA, destinada ao depósito da {bolsa} concedida pela UENF.')
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = Inches(0.25)

    doc.add_paragraph()
    doc.add_paragraph('Sem mais, agradecemos.')

    doc.add_paragraph()
    doc.add_paragraph()
    para = doc.add_picture('Assinatura.png')
    doc.paragraphs[18].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph('_______________________________________')
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph('Clicia Grativol Gaspar de Matos')
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph('Pró-Reitora de Assuntos Comunitários')
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    para = doc.add_paragraph()
    bold_para = para.add_run('Via Bradesco')
    bold_para.bold = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('_____________________________________________________________________________')
    para = doc.add_paragraph()
    doc.add_paragraph('Prezado(a) Solicitante,')
    doc.add_paragraph()
    para = doc.add_paragraph(f'Este protocolo de Abertura de Conta Corrente deverá ser preenchido e entregue à Pró-Reitoria de Assuntos Comunitários para cadastro e recebimento da {bolsa}.')
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = Inches(0.25)
    doc.add_paragraph()

    tabela = doc.add_table(rows=3, cols=2)
    tabela.border = 

    row = tabela.rows[0]
    row.height = Mm(8.0)

    row = tabela.rows[1]
    row.height = Mm(8.0)

    row = tabela.rows[2]
    row.height = Mm(8.0)

    row = tabela.rows[0].cells
    row[0].paragraphs[0].add_run('Agência').bold = True
    row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row = tabela.rows[1].cells
    row[0].paragraphs[0].add_run('Conta Corrente').bold = True
    row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row = tabela.rows[2].cells
    row[0].paragraphs[0].add_run('Carimbo e assinatura do Gerente').bold = True
    row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    para = doc.add_paragraph()
    bold_para = para.add_run('Via Solicitante')
    bold_para.bold = True
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(f'docs\{cpf_formatado} - {nome}.docx')


def gerar_carta_old(cpf, nome, bolsa):
    documento = Document('Oficio_Abertura_de_Conta.docx')

    data = str(datetime.today())
    data, _ = data.split(' ')
    ano, mes, dia = data.split('-')
    mes = meses[int(mes)-1]
    data = dia + ' de ' + mes + ' de ' + ano

    cpf_formatado = cpf[0:3] + '.' + cpf[3:6] + '.' + cpf[6:9] + '-' + cpf[-2:]
    nome = nome.upper()
    bolsa = bolsa.upper()

    referencias = {
        '«Data»': data,
        '«CPF»': cpf_formatado,
        '«Nome»': nome,
        '«Bolsa»': bolsa,
        '«Bolsa2»': bolsa,
    }

    for paragrafo in documento.paragraphs:
        for codigo in referencias:
            valor = referencias[codigo]
            paragrafo.text = paragrafo.text.replace(codigo, '<b>' + valor + "</b>")

    for paragrafo in documento.paragraphs:
        for r in paragrafo.runs:
            print(f'{r.text} - {r.style.name}')

    documento.save(f'docs\{cpf_formatado} - {nome}.docx')

if __name__ == '__main__':
    gerar_carta('95204130730', 'luiz batista de almeida', 'bolsa de apoio acadêmico')
